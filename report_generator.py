"""
Generate DDR as DOCX (python-docx) and PDF (reportlab).
"""

from __future__ import annotations

import json
import logging
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = PROJECT_ROOT / "outputs"


def _format_match_strength(obs: dict[str, Any]) -> str:
    pct = obs.get("confidence_percent")
    tier = obs.get("confidence_tier")
    if pct is not None and tier:
        return f"{pct}% — {tier}"
    c = obs.get("confidence")
    if isinstance(c, (int, float)):
        return f"{int(round(min(1.0, max(0.0, float(c))) * 100))}%"
    return "—"


def _ensure_output_dir() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


def build_docx(report: dict[str, Any], output_path: Path | None = None) -> bytes:
    """Build Word document; return bytes. Optionally write to output_path."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    title = doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    title.runs[0].font.size = Pt(18)

    doc.add_paragraph(
        "This report was generated using rule-based document analysis. "
        "It is not a substitute for professional engineering or licensed inspection."
    )
    doc.add_paragraph("")

    doc.add_heading("1. Property Issue Summary", level=1)
    doc.add_paragraph(report.get("property_issue_summary", "Not Available"))

    doc.add_heading("2. Area-wise Observations", level=1)
    observations = report.get("observations") or []
    if not observations:
        doc.add_paragraph("No observations extracted. Image: Not Available")
    for idx, obs in enumerate(observations, start=1):
        doc.add_heading(f"2.{idx} {obs.get('area', 'Not Available')} — {obs.get('issue', '')}", level=2)
        doc.add_paragraph(f"Description: {obs.get('description', 'Not Available')}")
        doc.add_paragraph(f"Thermal: {obs.get('thermal_observation', 'Not Available')}")
        doc.add_paragraph(f"Combined insight: {obs.get('combined_insight', 'Not Available')}")
        doc.add_paragraph(
            f"Severity: {obs.get('severity', 'Not Available')} | "
            f"Match strength: {_format_match_strength(obs)}"
        )
        doc.add_paragraph(f"Recommendation: {obs.get('recommendation', 'Not Available')}")
        img = obs.get("image_path")
        if img and Path(img).is_file():
            try:
                doc.add_picture(img, width=Inches(5.5))
            except Exception as e:
                logger.warning("DOCX image insert failed: %s", e)
                doc.add_paragraph("Image Not Available")
        else:
            doc.add_paragraph("Image Not Available")
        doc.add_paragraph("")

    doc.add_heading("3. Probable Root Cause", level=1)
    doc.add_paragraph(report.get("probable_root_cause", "Not Available"))

    sev = report.get("severity_assessment") or {}
    doc.add_heading("4. Severity Assessment", level=1)
    doc.add_paragraph(f"Overall: {sev.get('overall', 'Not Available')}")
    doc.add_paragraph(f"Reasoning: {sev.get('reasoning', 'Not Available')}")
    ce = report.get("confidence_explanation")
    if ce:
        doc.add_paragraph(str(ce))

    doc.add_heading("5. Recommended Actions", level=1)
    for a in report.get("recommended_actions") or ["Not Available"]:
        doc.add_paragraph(str(a), style="List Bullet")

    doc.add_heading("6. Additional Notes", level=1)
    doc.add_paragraph(report.get("additional_notes", "Not Available"))

    doc.add_heading("7. Missing / Unclear Information", level=1)
    for m in report.get("missing_or_unclear") or ["Not Available"]:
        doc.add_paragraph(str(m), style="List Bullet")

    conflicts = report.get("conflicts") or []
    if conflicts:
        doc.add_heading("8. Conflicts Between Reports", level=1)
        for c in conflicts:
            doc.add_paragraph(str(c), style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if output_path:
        output_path.write_bytes(data)
    return data


def build_pdf(report: dict[str, Any], output_path: Path | None = None) -> bytes:
    """Build PDF with reportlab; return bytes."""
    buf = BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=72,
        bottomMargin=54,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "BodyJustify",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    h0 = ParagraphStyle("H0", parent=styles["Heading1"], fontSize=16, spaceAfter=12)
    h1 = ParagraphStyle("H1", parent=styles["Heading2"], fontSize=13, spaceAfter=8, spaceBefore=10)
    h2 = ParagraphStyle("H2", parent=styles["Heading3"], fontSize=11, spaceAfter=6, spaceBefore=8)
    story: list[Any] = []

    story.append(Paragraph("Detailed Diagnostic Report (DDR)", h0))
    story.append(
        Paragraph(
            "<i>Rule-based analysis. Not a substitute for licensed professional inspection.</i>",
            body,
        )
    )
    story.append(Spacer(1, 0.15 * inch))

    def add_h1(t: str) -> None:
        story.append(Paragraph(t.replace("&", "&amp;"), h1))

    def add_p(t: str) -> None:
        safe = t.replace("&", "&amp;").replace("<", "&lt;")
        story.append(Paragraph(safe, body))

    add_h1("1. Property Issue Summary")
    add_p(str(report.get("property_issue_summary", "Not Available")))

    add_h1("2. Area-wise Observations")
    observations = report.get("observations") or []
    if not observations:
        add_p("No observations extracted. Image Not Available.")
    for idx, obs in enumerate(observations, start=1):
        area = str(obs.get("area", "Not Available")).replace("&", "&amp;")
        issue = str(obs.get("issue", "")).replace("&", "&amp;")
        story.append(Paragraph(f"2.{idx} {area} — {issue}", h2))
        add_p(f"Description: {obs.get('description', 'Not Available')}")
        add_p(f"Thermal: {obs.get('thermal_observation', 'Not Available')}")
        add_p(f"Combined insight: {obs.get('combined_insight', 'Not Available')}")
        add_p(
            f"Severity: {obs.get('severity', 'Not Available')} | "
            f"Match strength: {_format_match_strength(obs)}"
        )
        add_p(f"Recommendation: {obs.get('recommendation', 'Not Available')}")
        img_path = obs.get("image_path")
        if img_path and Path(img_path).is_file():
            try:
                story.append(Spacer(1, 0.08 * inch))
                story.append(RLImage(str(Path(img_path).resolve()), width=5 * inch))
            except Exception as e:
                logger.warning("PDF image insert failed: %s", e)
                add_p("Image Not Available")
        else:
            add_p("Image Not Available")
        story.append(Spacer(1, 0.12 * inch))

    add_h1("3. Probable Root Cause")
    add_p(str(report.get("probable_root_cause", "Not Available")))

    sev = report.get("severity_assessment") or {}
    add_h1("4. Severity Assessment")
    add_p(f"Overall: {sev.get('overall', 'Not Available')}")
    add_p(f"Reasoning: {sev.get('reasoning', 'Not Available')}")
    if report.get("confidence_explanation"):
        add_p(str(report["confidence_explanation"]))

    add_h1("5. Recommended Actions")
    for a in report.get("recommended_actions") or ["Not Available"]:
        story.append(Paragraph(f"• {str(a).replace('&', '&amp;')}", body))

    add_h1("6. Additional Notes")
    add_p(str(report.get("additional_notes", "Not Available")))

    add_h1("7. Missing / Unclear Information")
    for m in report.get("missing_or_unclear") or ["Not Available"]:
        story.append(Paragraph(f"• {str(m).replace('&', '&amp;')}", body))

    conflicts = report.get("conflicts") or []
    if conflicts:
        add_h1("8. Conflicts Between Reports")
        for c in conflicts:
            story.append(Paragraph(f"• {str(c).replace('&', '&amp;')}", body))

    pdf.build(story)
    data = buf.getvalue()
    if output_path:
        output_path.write_bytes(data)
    return data


def save_json_report(report: dict[str, Any], output_path: Path) -> None:
    """Write JSON without non-serializable paths if any."""
    clean = json.loads(json.dumps(report, default=str))
    output_path.write_text(json.dumps(clean, indent=2), encoding="utf-8")
